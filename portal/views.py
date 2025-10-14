from __future__ import annotations

import logging
import re
import shutil
import tempfile
from pathlib import Path

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files import File
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse

from .decorators import portal_menu_required
from .forms import (
    AccessApprovalForm,
    FolderForm,
    LibraryDocumentForm,
    OcrRequestForm,
    PdfToWordForm,
    PortalSettingsForm,
    SignUpForm,
    WordDocumentForm,
)
from .models import (
    LibraryFolder,
    OcrJob,
    PortalAccess,
    PortalSettings,
    StoredDocument,
    WordDocument,
)

log = logging.getLogger(__name__)


@login_required
def home(request):
    access = getattr(request.user, 'portal_access', None)
    return render(request, 'portal/home.html', {'access': access})


def signup(request):
    if request.user.is_authenticated:
        messages.info(request, 'Ești deja autentificat.')
        return redirect('portal:home')

    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            access, _ = PortalAccess.objects.get_or_create(user=user)
            access.status = PortalAccess.Status.PENDING
            access.save(update_fields=['status', 'updated_at'])
            messages.success(
                request,
                'Cont creat cu succes. Un administrator trebuie să aprobe accesul înainte de utilizare.',
            )
            return redirect('login')
    else:
        form = SignUpForm()

    return render(request, 'portal/signup.html', {'form': form})


@portal_menu_required('ocr')
def ocr_studio(request):
    form = OcrRequestForm(request.POST or None, request.FILES or None, user=request.user)

    if request.method == 'POST' and form.is_valid():
        pdf_file = form.cleaned_data['pdf_file']
        language_codes = form.cleaned_language_codes()
        options = form.selected_options()
        destination_folder = form.cleaned_data.get('destination_folder')

        job = OcrJob(
            user=request.user,
            language=language_codes,
            status=OcrJob.Status.PROCESSING,
            options=options,
            destination_folder=destination_folder,
        )
        job.source_file.save(pdf_file.name, pdf_file, save=False)
        job.save()

        try:
            _run_ocr(job)
            messages.success(request, 'Documentul a fost procesat cu succes cu OCRmyPDF.')
        except RuntimeError as exc:
            job.status = OcrJob.Status.FAILED
            job.error_message = str(exc)
            job.save(update_fields=['status', 'error_message', 'updated_at'])
            messages.error(request, f'Procesarea a eșuat: {exc}')

        return redirect('portal:ocr')

    jobs = OcrJob.objects.filter(user=request.user)[:25]
    settings_obj = PortalSettings.load()
    engine_key = settings_obj.ocr_engine or PortalSettings.OcrEngine.OCRMYPDF
    engine_label = settings_obj.get_ocr_engine_display()
    ocrmypdf_version = None
    if engine_key == PortalSettings.OcrEngine.OCRMYPDF:
        try:
            import ocrmypdf

            ocrmypdf_version = ocrmypdf.__version__
        except Exception:  # pragma: no cover - optional metadata
            ocrmypdf_version = None
    return render(
        request,
        'portal/ocr_studio.html',
        {
            'form': form,
            'jobs': jobs,
            'ocrmypdf_version': ocrmypdf_version,
            'engine_label': engine_label,
            'engine_key': engine_key,
            'docling_available': PortalSettings.docling_available(),
        },
    )


@portal_menu_required('libraries')
def libraries(request):
    folder_form = FolderForm(request.POST or None, user=request.user)
    document_form = LibraryDocumentForm(
        request.POST or None, request.FILES or None, user=request.user
    )

    if request.method == 'POST':
        if 'create_folder' in request.POST and folder_form.is_valid():
            folder = folder_form.save(commit=False)
            folder.user = request.user
            folder.save()
            messages.success(request, f'Folderul „{folder.name}” a fost creat.')
            return redirect('portal:libraries')

        if 'upload_document' in request.POST and document_form.is_valid():
            folder = document_form.cleaned_data['folder']
            if folder.user != request.user:
                messages.error(request, 'Nu poți încărca documente în foldere care nu îți aparțin.')
                return redirect('portal:libraries')
            document = document_form.save(commit=False)
            document.folder = folder
            document.save()
            document_form.save_m2m()
            messages.success(request, f'Documentul „{document.title}” a fost încărcat.')
            return redirect('portal:libraries')

    folders = (
        request.user.library_folders.select_related('parent').prefetch_related('documents')
    )
    return render(
        request,
        'portal/libraries.html',
        {
            'folders': folders,
            'folder_form': folder_form,
            'document_form': document_form,
        },
    )


@portal_menu_required('libraries')
def library_detail(request, folder_id):
    folder = get_object_or_404(LibraryFolder, id=folder_id, user=request.user)
    documents = folder.documents.select_related('ocr_job').all()
    return render(
        request,
        'portal/library_detail.html',
        {
            'folder': folder,
            'documents': documents,
        },
    )


@portal_menu_required('preview')
def preview_hub(request):
    documents = (
        StoredDocument.objects.filter(folder__user=request.user)
        .select_related('folder')
        .order_by('-updated_at')
    )
    return render(
        request,
        'portal/preview_hub.html',
        {
            'documents': documents,
        },
    )


@portal_menu_required('preview')
def preview_document(request, document_id):
    document = get_object_or_404(StoredDocument, id=document_id, folder__user=request.user)
    return render(
        request,
        'portal/preview_document.html',
        {
            'document': document,
        },
    )


@portal_menu_required('preview')
def download_document(request, document_id):
    document = get_object_or_404(StoredDocument, id=document_id, folder__user=request.user)
    if not document.processed_file:
        raise Http404('Documentul procesat nu este disponibil pentru descărcare.')
    return FileResponse(
        document.processed_file.open('rb'),
        as_attachment=True,
        filename=document.processed_filename() or 'document_ocr.pdf',
    )


@portal_menu_required('word')
def word_studio(request):
    create_form = WordDocumentForm(request.POST or None)
    convert_form = PdfToWordForm(request.POST or None, request.FILES or None)
    settings_obj = PortalSettings.load()

    if request.method == 'POST':
        if 'create_word' in request.POST and create_form.is_valid():
            title = create_form.cleaned_data['title']
            body = create_form.cleaned_data.get('body', '')
            word_document = _generate_docx(request.user, title, body)
            messages.success(request, f'Documentul Word „{word_document.title}” a fost creat.')
            return redirect('portal:word')

        if 'convert_pdf' in request.POST and convert_form.is_valid():
            pdf_file = convert_form.cleaned_data['pdf_file']
            title = convert_form.cleaned_data['title']
            try:
                word_document = _convert_pdf_to_word(request.user, title, pdf_file)
            except RuntimeError as exc:
                messages.error(request, f'Conversia a eșuat: {exc}')
            else:
                messages.success(
                    request,
                    f'PDF-ul a fost convertit cu succes în documentul Word „{word_document.title}”.',
                )
            return redirect('portal:word')

    documents = request.user.word_documents.all()
    return render(
        request,
        'portal/word_studio.html',
        {
            'create_form': create_form,
            'convert_form': convert_form,
            'documents': documents,
            'engine_label': settings_obj.get_ocr_engine_display(),
            'engine_key': settings_obj.ocr_engine,
        },
    )


@portal_menu_required('word')
def download_word_document(request, document_id):
    document = get_object_or_404(WordDocument, id=document_id, user=request.user)
    return FileResponse(
        document.document_file.open('rb'),
        as_attachment=True,
        filename=Path(document.document_file.name).name,
    )


@portal_menu_required('ocr')
def download_job(request, job_id):
    job = get_object_or_404(OcrJob, id=job_id, user=request.user)

    if job.status != OcrJob.Status.COMPLETED or not job.processed_file:
        raise Http404('Documentul nu este disponibil pentru descărcare.')

    return FileResponse(
        job.processed_file.open('rb'),
        as_attachment=True,
        filename=job.processed_filename() or 'document_ocr.pdf',
    )


@portal_menu_required('ocr')
def download_sidecar(request, job_id):
    job = get_object_or_404(OcrJob, id=job_id, user=request.user)
    if not job.sidecar_file:
        raise Http404('Fișierul sidecar nu este disponibil.')
    return FileResponse(
        job.sidecar_file.open('rb'),
        as_attachment=True,
        filename=job.sidecar_filename() or 'document.txt',
    )


@portal_menu_required('admin')
def admin_console(request):
    if not request.user.is_staff:
        messages.error(request, 'Doar administratorii pot accesa consola de administrare.')
        return redirect('portal:home')

    access_list = PortalAccess.objects.select_related('user').order_by('created_at')
    selected_id = request.GET.get('access_id')
    selected_access = None
    form = None
    settings_obj = PortalSettings.load()
    settings_form = PortalSettingsForm(
        request.POST if request.method == 'POST' and request.POST.get('form_name') == 'settings' else None,
        instance=settings_obj,
        prefix='settings',
    )

    if request.method == 'POST' and request.POST.get('form_name') == 'settings':
        settings_form = PortalSettingsForm(
            request.POST,
            instance=settings_obj,
            prefix='settings',
        )
        if settings_form.is_valid():
            settings_form.save()
            messages.success(
                request,
                f"Motorul OCR implicit a fost schimbat în {settings_form.instance.get_ocr_engine_display()}.",
            )
            return redirect('portal:admin')

    if selected_id:
        selected_access = get_object_or_404(PortalAccess, id=selected_id)
        if request.method == 'POST' and request.POST.get('form_name') == 'access':
            form = AccessApprovalForm(request.POST, instance=selected_access)
            if form.is_valid():
                access = form.save(commit=False)
                if access.status == PortalAccess.Status.APPROVED and not access.allowed_menus:
                    access.grant_defaults()
                access.save()
                messages.success(
                    request,
                    f'Drepturile pentru utilizatorul {access.user} au fost actualizate.',
                )
                return redirect(f"{reverse('portal:admin')}?access_id={access.id}")
        else:
            form = AccessApprovalForm(instance=selected_access)

    return render(
        request,
        'portal/admin_console.html',
        {
            'access_list': access_list,
            'selected_access': selected_access,
            'form': form,
            'settings_form': settings_form,
            'portal_settings': settings_obj,
        },
    )


def _run_ocr(job: OcrJob) -> None:
    settings_obj = PortalSettings.load()
    engine = settings_obj.ocr_engine or PortalSettings.OcrEngine.OCRMYPDF
    options = job.options or {}
    options['engine'] = engine
    job.options = options
    job.save(update_fields=['options'])

    if engine == PortalSettings.OcrEngine.DOCLING:
        _run_with_docling(job)
    else:
        _run_with_ocrmypdf(job)

    if job.destination_folder and job.status == OcrJob.Status.COMPLETED:
        stored = StoredDocument(
            folder=job.destination_folder,
            ocr_job=job,
            title=Path(job.source_file.name).stem,
        )
        with job.source_file.open('rb') as original_stream:
            stored.original_file.save(
                Path(job.source_file.name).name,
                File(original_stream),
                save=False,
            )
        if job.processed_file:
            with job.processed_file.open('rb') as processed_stream:
                stored.processed_file.save(
                    job.processed_filename(),
                    File(processed_stream),
                    save=False,
                )
        stored.save()


def _run_with_ocrmypdf(job: OcrJob) -> None:
    try:
        import ocrmypdf
        from ocrmypdf import exceptions as ocrmypdf_exceptions
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            'OCRmyPDF nu este instalat. Instalează pachetul „ocrmypdf” și dependențele Tesseract.'
        ) from exc

    job.ensure_directories()

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        input_path = temp_dir_path / 'input.pdf'
        output_path = temp_dir_path / 'output.pdf'
        sidecar_path = temp_dir_path / 'sidecar.txt'

        with job.source_file.open('rb') as uploaded, input_path.open('wb') as destination:
            shutil.copyfileobj(uploaded, destination)

        language = job.language or None
        options = job.options or {}
        if options.get('auto_language'):
            language = None

        ocr_kwargs = {
            'language': language,
            'optimize': int(options.get('optimize', 1) or 0),
            'deskew': options.get('deskew', False),
            'rotate_pages': options.get('rotate_pages', False),
            'remove_background': options.get('remove_background', False),
            'clean_final': options.get('clean_final', False),
            'skip_text': options.get('skip_text', False),
            'force_ocr': options.get('force_ocr', False),
            'output_type': options.get('output_type') or 'pdfa',
            'progress_bar': False,
        }

        sidecar_requested = options.get('make_sidecar')
        if sidecar_requested:
            ocr_kwargs['sidecar'] = str(sidecar_path)

        try:
            ocrmypdf.ocr(
                str(input_path),
                str(output_path),
                **ocr_kwargs,
            )
        except (
            ocrmypdf_exceptions.MissingDependencyError,
            ocrmypdf_exceptions.PriorOcrFoundError,
            getattr(ocrmypdf_exceptions, 'SubprocessOutputError', ocrmypdf_exceptions.OcrError),
            ocrmypdf_exceptions.OcrError,
        ) as exc:
            log.exception('OCR failed for job %s', job.id)
            raise RuntimeError(str(exc)) from exc

        with output_path.open('rb') as processed:
            job.processed_file.save(
                f"{Path(job.source_file.name).stem}_ocr.pdf",
                File(processed),
                save=False,
            )

        if sidecar_requested and sidecar_path.exists():
            with sidecar_path.open('rb') as sidecar_stream:
                job.sidecar_file.save(
                    f"{Path(job.source_file.name).stem}.txt",
                    File(sidecar_stream),
                    save=False,
                )
        elif job.sidecar_file:
            job.sidecar_file.delete(save=False)
            job.sidecar_file = None

    job.status = OcrJob.Status.COMPLETED
    job.error_message = ''
    job.save(
        update_fields=['processed_file', 'sidecar_file', 'status', 'error_message', 'updated_at']
    )


def _run_with_docling(job: OcrJob) -> None:
    try:
        from docling.document_converter import DocumentConverter
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            'Docling nu este instalat. Instalează pachetul „docling” pentru a folosi acest motor.'
        ) from exc

    job.ensure_directories()

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        input_path = temp_dir_path / 'input.pdf'
        output_path = temp_dir_path / 'output.pdf'
        sidecar_path = temp_dir_path / 'sidecar.txt'

        with job.source_file.open('rb') as uploaded, input_path.open('wb') as destination:
            shutil.copyfileobj(uploaded, destination)

        converter = DocumentConverter()
        result = converter.convert(str(input_path))
        document = getattr(result, 'document', None)
        if document is None:
            raise RuntimeError('Docling nu a putut procesa documentul furnizat.')

        pdf_bytes = getattr(result, 'pdf_bytes', None)
        if pdf_bytes:
            with output_path.open('wb') as pdf_out:
                pdf_out.write(pdf_bytes)
        elif hasattr(document, 'export_to_pdf'):
            exported_pdf = document.export_to_pdf()
            if isinstance(exported_pdf, (bytes, bytearray)):
                with output_path.open('wb') as pdf_out:
                    pdf_out.write(exported_pdf)
            else:
                shutil.copyfile(input_path, output_path)
        else:
            shutil.copyfile(input_path, output_path)

        text_content = ''
        if hasattr(document, 'export_to_markdown'):
            text_content = _markdown_to_plain_text(document.export_to_markdown())
        elif hasattr(document, 'export_to_text'):
            text_content = str(document.export_to_text())
        elif hasattr(document, 'pages'):
            lines = []
            for page in getattr(document, 'pages', []):
                page_text = getattr(page, 'text', '')
                if page_text:
                    lines.append(page_text)
            text_content = '\n'.join(lines)

        text_content = text_content.strip()
        if not text_content:
            text_content = 'Nu a fost posibilă extragerea textului cu Docling.'

        if options.get('make_sidecar'):
            sidecar_path.write_text(text_content, encoding='utf-8', errors='ignore')
            with sidecar_path.open('rb') as sidecar_stream:
                job.sidecar_file.save(
                    f"{Path(job.source_file.name).stem}.txt",
                    File(sidecar_stream),
                    save=False,
                )
        elif job.sidecar_file:
            job.sidecar_file.delete(save=False)
            job.sidecar_file = None

        with output_path.open('rb') as processed:
            job.processed_file.save(
                f"{Path(job.source_file.name).stem}_docling.pdf",
                File(processed),
                save=False,
            )

    job.status = OcrJob.Status.COMPLETED
    job.error_message = ''
    job.save(
        update_fields=['processed_file', 'sidecar_file', 'status', 'error_message', 'updated_at']
    )


def _markdown_to_plain_text(markdown_text: str) -> str:
    cleaned_lines = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append('')
            continue
        line = re.sub(r'^[#>*\-\d\.\s]+', '', line)
        line = line.replace('**', '').replace('*', '').replace('_', '')
        cleaned_lines.append(line.strip())
    return '\n'.join(cleaned_lines)


def _load_docx_document():
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            'Librăria python-docx nu este instalată. Adaugă „python-docx” în dependențe pentru a genera documente Word.'
        ) from exc
    return Document


def _generate_docx(user, title: str, body: str) -> WordDocument:
    Document = _load_docx_document()
    document = Document()
    document.add_heading(title, level=1)
    if body:
        for paragraph in body.split('\n'):
            document.add_paragraph(paragraph)

    with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
        document.save(tmp.name)
        tmp.seek(0)
        word_doc = WordDocument(user=user, title=title)
        word_doc.document_file.save(f"{title.replace(' ', '_')}.docx", File(tmp), save=False)
        word_doc.save()
    return word_doc


def _convert_pdf_to_word(user, title: str, pdf_file) -> WordDocument:
    Document = _load_docx_document()
    job = OcrJob(user=user, status=OcrJob.Status.PROCESSING, options={'make_sidecar': True})
    job.source_file.save(pdf_file.name, pdf_file, save=False)
    job.save()

    try:
        _run_ocr(job)
    except RuntimeError as exc:
        job.status = OcrJob.Status.FAILED
        job.error_message = str(exc)
        job.save(update_fields=['status', 'error_message', 'updated_at'])
        raise

    text_content = ''
    if job.sidecar_file:
        with job.sidecar_file.open('r', encoding='utf-8', errors='ignore') as sidecar:
            text_content = sidecar.read()

    document = Document()
    document.add_heading(title, level=1)
    for paragraph in text_content.split('\n'):
        if paragraph.strip():
            document.add_paragraph(paragraph)

    with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
        document.save(tmp.name)
        tmp.seek(0)
        word_doc = WordDocument(user=user, title=title)
        with job.source_file.open('rb') as source_stream:
            word_doc.source_pdf.save(Path(job.source_file.name).name, File(source_stream), save=False)
        word_doc.document_file.save(f"{title.replace(' ', '_')}.docx", File(tmp), save=False)
        word_doc.save()

    return word_doc
